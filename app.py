from flask import Flask, request, jsonify
from flask_cors import CORS
from PyPDF2 import PdfReader
import docx2txt
import fitz  # PyMuPDF
import os
import re
import io
import docx
import striprtf
import tempfile
import logging

app = Flask(__name__)
# السماح بجميع النطاقات مؤقتًا. في الإنتاج، حدد النطاقات المسموح بها بدقة.
CORS(app)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- إعدادات الأمان (هام: قم بتغيير هذا الرمز إلى قيمة قوية ومعقدة في بيئة الإنتاج!) ---
SECRET_ACCESS_CODE = "1234" # غير هذا الرمز إلى رمز آمن ومعقد!
# ----------------------------------------------------------------------------------

# استخراج النص من الملف
def extract_text(file_storage):
    filename = file_storage.filename.lower()
    file_storage.seek(0) # إعادة المؤشر للبداية قبل القراءة
    try:
        if filename.endswith(".pdf"):
            reader = PdfReader(file_storage)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text
        elif filename.endswith((".doc", ".docx")):
            # docx2txt يحتاج اسم ملف، نحفظ مؤقتا
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(file_storage.read())
                tmp_path = tmp.name
            try:
                text = docx2txt.process(tmp_path)
            finally:
                os.remove(tmp_path) # التأكد من حذف الملف المؤقت
            return text
        elif filename.endswith(".rtf"):
            # قد تحتاج لضبط التشفير حسب ملف RTF
            file_storage.seek(0)
            rtf_bytes = file_storage.read()
            rtf_text = rtf_bytes.decode("utf-8", errors="ignore") # محاولة UTF-8 أولاً
            try:
                return striprtf.rtf_to_text(rtf_text)
            except Exception:
                rtf_text = rtf_bytes.decode("cp1252", errors="ignore") # محاولة تشفير آخر
                return striprtf.rtf_to_text(rtf_text)
        elif filename.endswith(".txt"):
            file_storage.seek(0)
            return file_storage.read().decode("utf-8", "ignore")
        else:
            return ""
    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {e}", exc_info=True)
        return ""

# كشف الجداول والصور والخطوط وأحجامها في PDF وDOCX
def analyze_document_structure(file_storage):
    filename = file_storage.filename.lower()
    tables = 0
    images = 0
    fonts = set()
    font_sizes = [] # لتخزين أحجام الخطوط المكتشفة

    file_storage.seek(0) # إعادة المؤشر للبداية
    try:
        if filename.endswith(".pdf"):
            doc_data = file_storage.read()
            doc = fitz.open(stream=doc_data, filetype="pdf")
            for page in doc:
                images += len(page.get_images(full=True))
                # كشف جداول (تقريبي) - لا يزال تقريبيًا
                text_blocks = page.get_text("dict")
                # البحث عن أنماط الجداول: عدد كبير من الخطوط الرأسية والأفقية
                vert_lines = sum(1 for item in text_blocks.get("lines", []) if abs(item['bbox'][0] - item['bbox'][2]) < 2 and item['bbox'][3] - item['bbox'][1] > 10)
                horiz_lines = sum(1 for item in text_blocks.get("lines", []) if abs(item['bbox'][1] - item['bbox'][3]) < 2 and item['bbox'][2] - item['bbox'][0] > 10)
                if vert_lines > 5 and horiz_lines > 5: # عتبة تقديرية
                    tables += 1

                for b in page.get_text("dict")["blocks"]:
                    for l in b.get("lines", []):
                        for s in l.get("spans", []):
                            font_name = s.get("font", "").split('+')[-1] # إزالة البادئة العشوائية (إذا وجدت)
                            font_size = s.get("size", 0)
                            if font_name:
                                fonts.add(font_name)
                            if font_size > 0:
                                font_sizes.append(font_size)
            doc.close()

        elif filename.endswith(".docx"):
            doc = docx.Document(file_storage)
            tables = len(doc.tables)
            # استخدام related_parts لاكتشاف الصور (أفضل من inline_shapes على السيرفرات)
            images = sum(1 for rel in doc.part.related_parts if "image" in rel.id and "image" in rel.target_parts[0].content_type)

            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font.name:
                        fonts.add(run.font.name)
                    if run.font.size:
                        # font.size يعود بكائن Pt، نحوله إلى قيمة عددية
                        font_sizes.append(run.font.size.pt)
            for table in doc.tables: # فحص الخطوط داخل الجداول أيضًا
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.font.name:
                                    fonts.add(run.font.name)
                                if run.font.size:
                                    font_sizes.append(run.font.size.pt)

    except Exception as e:
        logger.error(f"Error detecting tables/images/fonts: {e}", exc_info=True)
    return tables, images, list(fonts), font_sizes

# تحليل الكلمات المفتاحية والأقسام لحساب النسبة الإجمالية
def analyze_text(text):
    if not text or len(text.strip()) < 100:
        # إذا كان النص قصيرًا جدًا أو فارغًا، أعد تقييمًا منخفضًا جدًا
        return 0, [], [], {
            "experience": 0, "education": 0, "skills": 0,
            "contact": 0, "objective_summary": 0, "certification": 0
        }

    # كلمات مفتاحية مع أوزان (تقريبية) لكل قسم
    keywords_weighted = {
        'experience': {
            'keywords': ['experience', 'work history', 'employment', 'responsibilities', 'achievements', 'projects', 'job title',
                         'خبرة', 'سجل وظيفي', 'توظيف', 'مهام', 'إنجازات', 'مشاريع', 'المسمى الوظيفي'],
            'weight': 30
        },
        'education': {
            'keywords': ['education', 'degree', 'university', 'bachelor', 'master', 'phd', 'college', 'graduation',
                         'دراسة', 'تعليم', 'شهادة', 'جامعة', 'بكالوريوس', 'ماجستير', 'دكتوراه', 'تخرج'],
            'weight': 20
        },
        'skills': {
            'keywords': ['skills', 'proficient', 'expertise', 'technologies', 'tools', 'languages', 'competencies',
                         'مهارات', 'إجادة', 'خبرة فنية', 'تقنيات', 'أدوات', 'لغات', 'كفاءات'],
            'weight': 25
        },
        'contact': {
            'keywords': ['email', 'phone', 'contact', 'linkedin', 'portfolio', 'github', 'address', 'website',
                         'بريد', 'هاتف', 'تواصل', 'لينكد إن', 'محفظة أعمال', 'عنوان', 'موقع الكتروني'],
            'weight': 10
        },
        'objective_summary': { # تم تغيير الاسم ليكون أكثر شمولاً
            'keywords': ['objective', 'summary', 'profile', 'career goal', 'about me', 'introduction',
                         'هدف', 'ملخص', 'نبذة', 'نبذة عني', 'ملف شخصي', 'مقدمة'],
            'weight': 10
        },
        'certification': {
            'keywords': ['certification', 'certified', 'license', 'licensure', 'award', 'workshop',
                         'شهادة', 'اعتماد', 'رخصة', 'جائزة', 'ورشة عمل'],
            'weight': 5
        }
    }

    section_scores = {}
    found_sections = []
    missing_sections = []
    total_score = 0
    text_lower = text.lower()

    for key, data in keywords_weighted.items():
        present = any(v in text_lower for v in data['keywords'])
        if present:
            section_scores[key] = 100 # ما زال تقييم 100% إذا وجد
            found_sections.append(key)
            total_score += data['weight']
        else:
            section_scores[key] = 0
            missing_sections.append(key)
            # لا نخصم، بل نبني النتيجة من الأقسام الموجودة

    # عقوبة الرموز غير التقليدية (الرموز التي قد تسبب مشاكل في أنظمة ATS القديمة)
    # هذه الرموز تشمل بعض أشكال النقاط التعدادية غير القياسية والسهام.
    # يتم الخصم فقط إذا لم تكن السيرة الذاتية طويلة بما يكفي لتوقع استخدامها في أقسام مثل المهارات أو الخبرة.
    if re.search(r'[\u2022\u25CF\u25BA\u25BC\u25B6\u25C0\u25C6\u25A0\u2713\u2714\u2715\u2716\u2705\u2192\u2190\u2191\u2193]', text_lower) and len(text) < 500:
        total_score -= 5
        logger.info("Found potentially problematic symbols in short resume, deducting 5 points.")

    # عقوبة الطول الزائد (قد يكون مبالغاً فيه وغير مفضل للـ ATS)
    if len(text.split()) > 700: # عدد الكلمات بدلاً من الحروف
        total_score -= 5
        logger.info("Resume too long (over 700 words), deducting 5 points.")
    elif len(text.split()) < 150: # عقوبة للسيرة الذاتية القصيرة جداً
        total_score -= 10
        logger.info("Resume too short (under 150 words), deducting 10 points.")


    # ضمان أن النتيجة لا تقل عن صفر وتتجاوز 100
    return min(100, max(0, total_score)), found_sections, missing_sections, section_scores

# تحليل التوصيات بناءً على الخط والحجم والجداول/الصور
def suggest_improvements(fonts, font_sizes, tables, images):
    suggestions = set()
    notes = []
    # قائمة أوسع من الخطوط الآمنة والمتوافقة مع ATS
    ats_safe_fonts = {"Arial", "Calibri", "Times New Roman", "Georgia", "Helvetica", "Verdana", "Roboto", "Lato", "Open Sans", "Tahoma", "Garamond"}

    # تحليل الخطوط المستخدمة
    if fonts:
        for font in fonts:
            # التحقق من الخطوط بأسماء جزئية أيضًا (مثل ArialMT -> Arial)
            found_safe_font = False
            for safe_font in ats_safe_fonts:
                if safe_font.lower() in font.lower():
                    found_safe_font = True
                    break
            
            if not found_safe_font:
                notes.append(f"الخط '{font}' قد لا يكون مدعومًا بشكل كامل في أنظمة ATS القديمة.")
                suggestions.add("استخدم خطوطًا شائعة مثل Arial, Calibri, Times New Roman, Helvetica لضمان أقصى توافق.")
    else:
        notes.append("لم يتمكن النظام من تحديد الخطوط المستخدمة. قد يكون هذا بسبب تنسيق الملف أو محتواه. يفضل استخدام خطوط نصية قياسية.")


    # تحليل أحجام الخطوط
    if font_sizes:
        # حساب متوسط حجم الخطوط أو حجم الخط الأكثر تكرارًا (المتوسط هنا)
        avg_font_size = sum(font_sizes) / len(font_sizes)
        
        # تحويل الأحجام إلى أرقام صحيحة أو عشرية لسهولة المقارنة
        if avg_font_size < 9.5: # أقل من 9.5 نقطة يعتبر صغيراً جداً
            notes.append(f"متوسط حجم الخط هو {avg_font_size:.1f} نقطة، وهو صغير جداً للقراءة في أنظمة ATS وربما للبشر.")
            suggestions.add("اجعل حجم الخط الرئيسي في السيرة الذاتية بين 10 و 12 نقطة.")
        elif avg_font_size > 14.5: # أكبر من 14.5 نقطة يعتبر كبيراً جداً
            notes.append(f"متوسط حجم الخط هو {avg_font_size:.1f} نقطة، وهو كبير جداً وقد يجعل السيرة الذاتية تبدو غير احترافية.")
            suggestions.add("اجعل حجم الخط الرئيسي في السيرة الذاتية بين 10 و 12 نقطة.")
        elif avg_font_size < 10:
             suggestions.add("حجم الخط يبدو صغيراً بعض الشيء (أقل من 10 نقاط). قد يكون من الأفضل زيادته قليلاً لضمان سهولة القراءة.")
        elif avg_font_size > 12:
             suggestions.add("حجم الخط يبدو كبيراً بعض الشيء (أكثر من 12 نقطة). قد يكون من الأفضل تقليله قليلاً للحفاظ على الإيجاز.")


    # ملاحظات على الجداول والصور
    if tables > 0:
        notes.append(f"تم اكتشاف {tables} جدول في الملف، وهذا قد يسبب مشاكل لبعض أنظمة ATS في قراءة المحتوى بشكل صحيح.")
        suggestions.add("تجنب استخدام الجداول قدر الإمكان في السيرة الذاتية، أو قم بتحويل محتواها إلى نص عادي.")
    if images > 0:
        notes.append(f"تم اكتشاف {images} صورة/عنصر رسومي في الملف، ويفضل تجنب الصور تمامًا في السيرة الذاتية.")
        suggestions.add("يفضل عدم تضمين الصور أو العناصر الرسومية (مثل الرسوم البيانية) في السيرة الذاتية الموجهة لأنظمة ATS.")
    
    return notes, list(suggestions)

@app.route('/analyze', methods=['POST'])
def analyze():
    try:
        if 'resume' not in request.files:
            return jsonify({'error': 'لم يتم رفع ملف السيرة الذاتية'}), 400
        uploaded_file = request.files['resume']
        if uploaded_file.filename == '':
            return jsonify({'error': 'لم يتم اختيار ملف'}), 400

        # حفظ الملف في ذاكرة مؤقتة لتمريره لدوال متعددة
        file_content = uploaded_file.read()
        file_stream = io.BytesIO(file_content)
        file_stream.filename = uploaded_file.filename # للحفاظ على اسم الملف

        text = extract_text(file_stream)
        if not text or len(text.strip()) < 10:
            logger.error("No text extracted from file!")
            return jsonify({'error': 'فشل في استخراج النص من الملف أو الملف فارغ أو غير مدعوم. يرجى التأكد من أن الملف نصي وقابل للقراءة.'}), 400

        # إعادة تهيئة stream بعد استخراج النص (لا تزال ضرورية)
        file_stream.seek(0)
        tables, images, fonts_used, font_sizes = analyze_document_structure(file_stream)

        # تحليل النص
        score, found, missing, section_scores = analyze_text(text)

        # اقتراحات بناءً على التحليل الهيكلي والتنسيقي
        notes, suggestions = suggest_improvements(fonts_used, font_sizes, tables, images)

        # تحليل تطابق وصف الوظيفة (مطابقة الكلمات)
        job_description = request.form.get('job_description', '').lower()
        match_score = None
        if job_description:
            # تنظيف النص من الرموز وعلامات الترقيم لتحسين المطابقة
            text_lower = text.lower() # تأكد من أن النص هنا lowercase
            cleaned_resume_text = re.sub(r'[^\w\s]', '', text_lower)
            cleaned_jd_text = re.sub(r'[^\w\s]', '', job_description)

            resume_words = set(cleaned_resume_text.split())
            jd_words = set(cleaned_jd_text.split())

            if jd_words:
                common = resume_words.intersection(jd_words)
                # حساب نسبة الكلمات المشتركة بناءً على كلمات وصف الوظيفة
                match_score = round((len(common) / len(jd_words)) * 100)
            else:
                match_score = 0
            
            # ملاحظة: لتحسين دقة match_score بشكل كبير، ستحتاج إلى NLP أعمق
            # مثل استخدام نماذج الكلمات (word embeddings) لحساب التشابه الدلالي.

        return jsonify({
            'score': score,
            'details': {
                'found': found,
                'missing': missing,
                'sections': section_scores
            },
            'fonts': list(fonts_used), # تحويل المجموعة إلى قائمة
            'font_sizes_detected': [round(s, 1) for s in font_sizes], # أحجام الخطوط الفعلية المكتشفة
            'notes': notes,
            'suggestions': suggestions,
            'match_score': match_score,
            'tables': tables,
            'images': images
        })
    except Exception as e:
        logger.error(f"Exception in analyze: {e}", exc_info=True)
        return jsonify({'error': 'خطأ غير متوقع أثناء تحليل السيرة الذاتية', 'details': str(e)}), 500

# نقطة نهاية للتحقق من الرمز السري (تستخدمها الواجهة الأمامية)
@app.route('/check_code', methods=['POST'])
def check_code():
    data = request.get_json()
    code = data.get('code')
    if code == SECRET_ACCESS_CODE:
        return jsonify({'success': True}), 200
    else:
        return jsonify({'success': False, 'message': 'رمز سري غير صحيح'}), 401 # Unauthorized

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({
        'status': 'healthy'
    })

if __name__ == '__main__':
    # استخدم debug=True فقط في بيئة التطوير، اجعله False في الإنتاج
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
